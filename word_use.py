# generate_test_doc.py
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_test_document():
    doc = Document()

    # ========== 标题层级 ==========
    doc.add_heading('一级标题：Word 转 Markdown 测试文档', level=1)
    doc.add_paragraph('这是一个用于测试转换功能的文档，包含常见的文档元素。')

    doc.add_heading('1. 列表测试', level=2)
    doc.add_paragraph('无序列表：', style='List Bullet')
    doc.add_paragraph('支持多级列表', style='List Bullet')
    doc.add_paragraph('可以包含粗体、斜体、代码等', style='List Bullet')

    doc.add_paragraph('有序列表：', style='List Number')
    doc.add_paragraph('第一项', style='List Number')
    doc.add_paragraph('第二项', style='List Number')
    doc.add_paragraph('第三项', style='List Number')

    # ========== 文本样式 ==========
    doc.add_heading('2. 文本样式测试', level=2)
    p = doc.add_paragraph()
    p.add_run('这是普通文本。').bold = False
    p.add_run('这是粗体文本。').bold = True
    p.add_run('这是斜体文本。').italic = True
    p.add_run('这是粗斜体。').bold = True
    p.add_run('这是').italic = True
    p.add_run(' 内联代码 ').font.name = 'Consolas'

    # ========== 引用 ==========
    doc.add_heading('3. 引用与代码块', level=2)
    doc.add_paragraph('这是引用块：', style='Quote')
    doc.add_paragraph('人生苦短，我用 Python。', style='Quote')
    
    # 模拟代码块（用等宽字体）111
    code_para = doc.add_paragraph()
    code_run = code_para.add_run('def hello():\n    print("Hello, World!")')
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(10)
    code_para.paragraph_format.left_indent = Inches(0.5)

    # ========== 表格 ==========
    doc.add_heading('4. 表格测试', level=2)
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    # 表头
    header_cells = table.rows[0].cells
    header_cells[0].text = '姓名'
    header_cells[1].text = '年龄'
    header_cells[2].text = '城市'
    header_cells[3].text = '备注'
    # 数据行
    data = [('张三', 28, '北京', '工程师'),
            ('李四', 32, '上海', '设计师'),
            ('王五', 25, '深圳', '产品经理')]
    for i, row_data in enumerate(data, start=1):
        cells = table.rows[i].cells
        for j, value in enumerate(row_data):
            cells[j].text = str(value)

    # ========== 图片（如果存在） ==========
    doc.add_heading('5. 图片测试', level=2)
    # 如果本地有图片可以添加，这里假设有 logo.png
    try:
        doc.add_picture('test_image.png', width=Inches(2))
        doc.add_paragraph('这是一张示例图片。')
    except:
        doc.add_paragraph('（图片占位符：请将 test_image.png 放在同目录下）')

    # ========== 链接与超链接 ==========
    doc.add_heading('6. 超链接测试', level=2)
    p = doc.add_paragraph()
    p.add_run('访问 GitHub: ')

    # ========== 特殊字符 ==========
    doc.add_heading('7. 特殊字符测试', level=2)
    doc.add_paragraph('包含中文：中文测试内容。')
    doc.add_paragraph('包含Emoji：😊 🎉 🔥')
    doc.add_paragraph('包含HTML实体：&lt; &gt; &amp;')

    # ========== 保存文档 ==========
    filename = 'test_document.docx'
    doc.save(filename)
    print(f'[SUCCESS] 测试文档已生成: {filename}')

if __name__ == '__main__':
    create_test_document()